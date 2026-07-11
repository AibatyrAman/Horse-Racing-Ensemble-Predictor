"""
KollektifÖğrenme-Ganyan-v2.docx üretici.

TNKÜ FBE tez makalesi formatını (KollektifÖğrenme-Ganyan.docx) referans alarak,
güncel/sızıntısız (TimeSeriesSplit) model sonuçlarını + piyasa ablasyonu +
kalibrasyon + McNemar bahis-kenarı analizini içeren genişletilmiş bildiri taslağını
üretir. Model karşılaştırma tabloları reports/model_comparison*.csv'den canlı okunur;
kalibrasyon/McNemar/kasa sayıları reports/calibration_summary.md ve
reports/betting_strategy_summary.md kaynaklıdır (bkz. docs/RAPOR_FAZ3_BAHIS_STRATEJISI.md).

Çıktı: proje kök dizininde KollektifÖğrenme-Ganyan-v2.docx (orijinal dosyaya dokunulmaz).
"""
import os

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(BASE_DIR)
OUT_PATH = os.path.join(ROOT, "KollektifÖğrenme-Ganyan-v2.docx")

FONT = "Times New Roman"
BODY_SIZE = Pt(12)
TABLE_SIZE = Pt(10)

TITLE = ("Kolektif Öğrenme Modelleriyle At Yarışı Tahmini: Sızıntısız Doğrulama, "
          "Piyasa Sinyali Ablasyonu ve İstatistiksel Bahis Kenarı Analizi")

KEYWORDS = ["Yapay zeka", "Topluluk modelleri", "At yarışı tahmini",
            "Zaman-serisi çapraz doğrulama", "Olasılık kalibrasyonu",
            "McNemar testi", "LightGBM", "XGBoost"]

ABSTRACT = (
    "Yapay zeka teknolojileri spor veri analitiğinde önemli ilerlemeler sağlamasına rağmen, "
    "at yarışı tahmininde sızıntısız (leakage-free) doğrulama ve istatistiksel anlamlılık "
    "testleriyle desteklenmiş çok boyutlu çalışmalar sınırlı kalmıştır. Bu çalışmada, "
    "Türkiye Jokey Kulübü (TJK) bünyesindeki profesyonel yerli at yarışlarında birinciyi "
    "(Is_Winner) ve ilk üçü (Is_Top3) tahmin etmek amacıyla 9 farklı makine öğrenmesi/"
    "topluluk modelinin performansı, ileri-zincirli zaman-serisi çapraz doğrulaması "
    "(TimeSeriesSplit) ile sızıntısız biçimde karşılaştırılmıştır. 21.648 yerli yarış kaydı "
    "(2.199 yarış, 127 gün) üzerinde Optuna ile optimize edilen modellerden, birinci "
    "tahmininde CatBoost en yüksek ayırt ediciliğe (AUC=0.8412), StackingEnsemble en yüksek "
    "pratik getiriye (P@1=%39.9, Değer Bahsi ROI=+%86.8) ulaşmıştır; tabela tahmininde ise "
    "XGBoost (AUC=0.8158, P@1=%73.2, P@3=%96.8) öne çıkmıştır. Piyasa sinyali (Ganyan oranı) "
    "çıkarıldığında AUC'deki ortalama düşüş yalnızca 0.0137-0.0222 olup, modelin gücünün "
    "büyük ölçüde piyasadan bağımsız özelliklerden geldiğini göstermiştir. Modelin piyasaya "
    "karşı seçim isabeti, McNemar testiyle altı bahis türünün tümünde istatistiksel olarak "
    "anlamlı bulunmuştur (Ganyan: +8.1 puan, p=1.6×10⁻¹⁶). Kalibrasyon analizi, kazanan "
    "olasılıklarının neredeyse kusursuz (ECE=0.0054), tabela olasılıklarının ise "
    "sınıf-dengeleme kaynaklı aşırı güvenli (ECE=0.3018) olduğunu ortaya koymuştur."
)

GIRIS_INTRO = (
    "Son yıllarda, makine öğrenmesi (ML) ve yapay zeka (AI) alanındaki araştırmalar kayda "
    "değer ilerlemeler kaydetmiştir. AI teknolojisi kullanılarak gerçekleştirilen "
    "sınıflandırma ve tahmin görevleri tıp, finans ve spor analitiği alanlarında derin "
    "etkiler yaratmıştır. Hipodrom yarışları ve atçılık sporu da bu alanlardan biridir. "
    "Atlı sporlarda, yarış sonuçlarını yüksek doğruluk oranlarıyla tahmin etmek, genetik "
    "kalıtım, anlık sağlık durumu, jokey deneyimi ve çevresel şartlar gibi karmaşık "
    "etkileşimlerin modellenmesini gerektirir. Bu durum klasik istatistik ve el yordamıyla "
    "hesaplanan handikap puanlamalarından çok daha karmaşık bir örüntü tanıma (pattern "
    "recognition) problemidir. Literatürde yapay zekanın veterinerlik tıbbına ve at "
    "sağlığına uygulanmasında (örneğin kolik hastalığının veya ameliyat sonrası hayatta "
    "kalma olasılığının tespiti) büyük başarılar elde edilmiş olsa da, çok boyutlu yarış "
    "verilerindeki makine öğrenmesi uygulamalarına yönelik derinlemesine analiz odaklı "
    "çalışmalar kısıtlıdır."
)

LITERATUR_OZETI = (
    "At ve yarış tahminleriyle ilgili literatür incelendiğinde, araştırmacıların "
    "çoğunlukla regresyon tabanlı tahmin yöntemleri veya daha sınırlı veri setleriyle "
    "çalışan düşük hacimli yapay sinir ağları üzerine odaklandıkları görülmektedir. Yurt "
    "içi tıp ve veterinerlik uygulamalarında, Cihan (2024) [1], kolikli atların cerrahi "
    "gerekliliğini makine öğrenmesi ile %76 ile %85 arasında bir doğruluk oranında tahmin "
    "etmeyi başarmıştır. Bunun dışında yarış süresini ve sıralamasını tahmin etmeyi "
    "hedefleyen geleneksel istatistiksel modeller, atın salt yaşını, kilosunu veya pist "
    "türünü temel almıştır. Literatürün geneli klasik makine öğrenmesi (Destek Vektör "
    "Makineleri, K-En Yakın Komşu vb.) üzerine kurgulanmış olup; aşırı uyum (overfitting) "
    "ihtimalini düşüren gelişmiş Boosting algoritmalarının (LightGBM, XGBoost, CatBoost) "
    "yarış dinamiği analizinde ve bilhassa idman sprint performanslarının zamansal "
    "başarıya (Time-Based Cross Validation) etkisinin değerlendirilmesinde önemli bir "
    "boşluk bulunmaktadır."
)

AMAC_KAPSAM = (
    "Bu çalışma, TJK (Türkiye Jokey Kulübü) veritabanında yer alan profesyonel yerli at "
    "yarışlarında birinciyi (Is_Winner) ve tabelaya ilk 3 bağlamında girmeyi (Is_Top3) "
    "tahmin etmek üzere makine öğrenmesi temelli, sızıntısız (leakage-free) bir altyapı "
    "oluşturmayı hedeflemektedir. Bu amaç doğrultusunda 1 yıllık süre zarfında kaydedilen "
    "21.648 veri satırından (2.199 yarış, 127 yarış günü) oluşan bir set işlenmiştir. "
    "Çalışmanın temel amacı; atların idman dereceleri, geçmiş anne-baba kazanma oranları "
    "(sızıntısız hedef kodlamasıyla hesaplanmış), göreceli sıklet ve göreceli handikap "
    "puanları gibi özellik mühendisliği yöntemleriyle yapılandırılmış setler üzerinde 9 "
    "farklı makine öğrenmesi/topluluk modelinin performans karşılaştırmasını sunmaktır. "
    "Çalışma ayrıca üç ek doğrulama katmanı eklemektedir: (i) piyasa sinyalinin (Ganyan "
    "oranı) modele katkısını ayrıştırmak için ablasyon analizi, (ii) tahmin "
    "olasılıklarının gerçek frekanslarla örtüşüp örtüşmediğini ölçen kalibrasyon analizi "
    "(Brier/ECE), ve (iii) modelin piyasaya karşı bahis seçimi isabetinin şans eseri olup "
    "olmadığını test eden McNemar anlamlılık testi. Bu sayede genel bahis kitlelerinin "
    "algılarından (Ganyan Baseline) kaynaklı kör bahis risklerinin giderilmesi ve doğruluk "
    "payı yüksek, istatistiksel olarak savunulabilir karar mekanizmalarının tesis "
    "edilmesi öngörülmektedir."
)

VERI_SETI = (
    "Çalışmanın ham verisi, bağımsız bir şekilde geliştirilen çok aşamalı veri kazıma "
    "(Web Scraping) otomasyonları vasıtasıyla derlenmiştir. Nisan 2025 ile Mart 2026 "
    "tarihleri arasında gerçekleşen koşu verileri taranmış, idman verisine ulaşılamayan "
    "dış yarışlar filtrelenmiştir. Sonuç itibarıyla tamamı yerli atların bulunduğu 2.199 "
    "yarışlık, yüksek sinyal-gürültü oranına sahip 21.648 veri noktalı bir ana veri setine "
    "erişilmiştir. Kategorik değişkenler (hedef sızıntısı - data leakage engellenerek) "
    "geçmiş kümülatif başarılara oranlanmış, aykırı veya NaN barındıran satırlar sistemli "
    "bir şekilde impute edilmiş (doldurulmuş) veya atılmıştır."
)

TS_SPLIT = (
    "Modellerin değerlendirilmesinde, yarışların kronolojik sırası korunarak "
    "ileri-zincirli zaman-serisi çapraz doğrulaması (TimeSeriesSplit) kullanılmıştır; her "
    "fold yalnızca geçmiş yarışlarla eğitilip gelecekteki yarışlarla test edilmiştir. Aynı "
    "yarışa ait tüm atlar aynı fold'a atanarak yarış bütünlüğü korunmuş, eksik-değer "
    "doldurma (imputation) ve kategorik kodlama adımları her fold'da yalnızca eğitim "
    "kısmı üzerinde fit edilmiştir; bu sayede test foldlarına ait istatistiklerin eğitime "
    "sızması önlenmiştir. Stacking Ensemble modelinin meta-özellik üretiminde kullanılan "
    "iç çapraz doğrulama da aynı ilkeyle TimeSeriesSplit'e dayandırılmıştır."
)

MODEL_ALGO = (
    "Model sınıflandırma mimarisinde çeşitli ekollerden toplam 9 adet model "
    "değerlendirilmiştir. Doğrusal ve ağaç temelli (Logistic Regression, Random Forest), "
    "Gradyan Arttırılmış Ağaçlar (GBM'ler; viz., Gradient Boosting, XGBoost, LightGBM, "
    "CatBoost) ve son olarak Topluluk (Ensemble) algoritmaları (Voting Classifier, "
    "Bagging, Stacking Classifier) analiz edilmiştir. Hiperparametre optimizasyon süreci, "
    "50 denemeli \"Optuna\" çatısı kullanılarak Eğri Altında Kalan Alan (AUC) skorunu "
    "maksimize etme üzerine kurgulanmıştır."
)

MODEL_ENSEMBLE_DETAY = (
    "Çalışmada oluşturulan topluluk algoritmalarının veri mimarisi özel olarak "
    "tasarlanmıştır. Voting Ensemble modeli, soft-voting yöntemiyle tahmin olasılıklarını "
    "ağırlıklı olarak birleştirmektedir; bu bağlamda sistemin işlem yükünü taşıyan "
    "boosting modellerine (LightGBM, XGBoost, CatBoost) 3 (üç) katsayı ağırlığı, Random "
    "Forest algoritmasına ise 1 (bir) katsayı ağırlığı atanmıştır. Daha sofistike bir "
    "birleştirici olan Stacking Ensemble modeli ise iki katmanlı bir yapıdan "
    "oluşmaktadır. Seviye-1 (ilk katman zayıf öğreniciler) tahmincileri olarak XGBoost, "
    "LightGBM, CatBoost, Random Forest ve veri standardizasyonu sağlanmış Lojistik "
    "Regresyon modelleri çalıştırılır. İlk katmandan süzülen tahmin olasılıkları "
    "meta-özelliklere dönüştürülür ve Seviye-2 nihai karar (meta-model) katmanına "
    "sunulur. Birinci katmanın çıktılarını ezberleyerek (overfitting) oluşabilecek "
    "zaafiyetleri önlemek maksadıyla meta-model olarak oldukça sert L2 regülarizasyonuna "
    "(C=1.0) sahip karmaşıklıktan uzak, doğrusal Lojistik Regresyon kullanılmıştır."
)

ABLASYON_YONTEM = (
    "Modelin başarısının piyasa oranına (Ganyan_Sayi) ne derece bağımlı olduğunu sınamak "
    "için, aynı pipeline bu değişken ve türevleri çıkarılarak yeniden çalıştırılmış; tam "
    "model ile piyasa-sinyalsiz (ablasyon) model AUC ve Precision@1 düzeyinde "
    "karşılaştırılmıştır. ΔAUC'nin küçük çıkması, modelin gücünün büyük ölçüde piyasa "
    "fiyatlamasını taklit etmekten değil, projeye özgü özellik mühendisliğinden "
    "(handikap, jokey/antrenör, soy hattı, idman) kaynaklandığına işaret eder."
)

KALIBRASYON_YONTEM = (
    "Modelin ürettiği olasılıkların gerçek isabet frekanslarıyla örtüşüp örtüşmediği, "
    "sızıntısız OOF (out-of-fold) tahminleri üzerinden Brier skoru ve Beklenen "
    "Kalibrasyon Hatası (Expected Calibration Error, ECE; 10 eşit-genişlikte bin) ile "
    "ölçülmüş, güvenilirlik diyagramlarıyla (reliability diagram) görselleştirilmiştir."
)

BAHIS_YONTEM = (
    "Tekil at kazanma olasılıklarından egzotik bahis türlerinin (İkili, Sıralı İkili, "
    "Üçlü, Tabela) kombinasyon olasılıkları Harville/Plackett-Luce modeliyle türetilmiş "
    "[9]; Harville formülünün favori atları sistematik olarak şişirme eğilimi λ=0.85 "
    "düzeltme katsayısıyla dengelenmiştir. Piyasaya kıyasla pozitif beklenen değer (EV) "
    "taşıyan bahisler, takas (takeout) sonrası +%5 EV eşiğini aşanlar olarak (Benter "
    "yaklaşımı [10]) filtrelenmiş; kasa yönetimi flat (%1 sabit pay) ve kesirli Kelly "
    "(¼-Kelly, %5 üst sınır) stratejileriyle paralel simüle edilmiştir. Modelin doğal "
    "seçiminin (olasılığa göre ilk-k) piyasanın doğal seçiminden (favori sırası) "
    "istatistiksel olarak anlamlı şekilde farklı olup olmadığı, eşleştirilmiş ikili "
    "sonuçlar için McNemar exact testi [11] ile sınanmıştır."
)

BULGU_WINNER = (
    "Baseline (referans) stratejisinde, piyasadaki genel kanının tespit ettiği en "
    "popüler atın yarışı kazanma ihtimali (P@1 - Doğruluk) %34.3, bahis bazında karşılığı "
    "ise -%26.2 (yatırım zararı) olarak ölçülmüştür. Sızıntısız TimeSeriesSplit "
    "değerlendirmesiyle elde edilen makine öğrenmesi sonuçları bu referansın belirgin "
    "biçimde üzerindedir (Tablo 3). En yüksek ayırt edicilik (AUC) skoruna CatBoost "
    "(0.8412) ulaşmış; nihai getiri ve sıra-1 isabet dengesinde ise StackingEnsemble "
    "(P@1=%39.9, Değer Bahsi ROI=+%86.8) üretim (production) modeli olarak belirlenmiştir."
)

BULGU_TOP3 = (
    "Hedef değişkenin ilk üç sıra olarak ele alındığı model kümesinde isabet oranları "
    "belirgin biçimde yükselmiştir. Tablo 4'te görüldüğü üzere XGBoost modeli "
    "(AUC=0.8158, P@1=%73.2, P@3=%96.8) bu hedefin öncüsü konumundadır. Veride yalnızca "
    "kazanma (Ganyan) ödemesi bulunduğundan, plase/tabela finişine parasal getiri (ROI) "
    "atfetmek yanıltıcı olacağından bu hedef için ROI raporlanmamış, değerlendirme "
    "sıralama metrikleriyle (AUC, P@1, P@3) sınırlandırılmıştır."
)

BULGU_ABLASYON = (
    "Ganyan_Sayi değişkeni ve türevleri çıkarıldığında, Is_Winner hedefinde ortalama AUC "
    "düşüşü +0.0137, Is_Top3 hedefinde +0.0222 olarak ölçülmüştür (Tablo 5-6, Şekil 3-4). "
    "Düşüşün küçüklüğü, modelin ayırt ediciliğinin büyük ölçüde piyasa fiyatlamasından "
    "değil; handikap, jokey/antrenör geçmişi, soy hattı ve idman verilerinden "
    "kaynaklandığını göstermektedir. Bu bulgu, modelin piyasa oranlarını yalnızca "
    "ezberlediği eleştirisini büyük ölçüde çürütmektedir."
)

BULGU_KALIBRASYON = (
    "17.923 at-koşu kaydı üzerindeki OOF olasılıkları incelendiğinde, Is_Winner "
    "tahminlerinin neredeyse kusursuz kalibre olduğu (Brier=0.0761, ECE=0.0054) "
    "görülmüştür. Buna karşın Is_Top3 tahminleri, üretim modelinde kullanılan "
    "sınıf-dengeleme (class-balancing) mekanizmaları nedeniyle yüksek-olasılık "
    "bölgesinde aşırı güvenli (over-confident) çıkmıştır (Brier=0.2630, ECE=0.3018); "
    "örneğin model %93.5 dediğinde gerçek isabet oranı yalnızca %71.2'dir (Tablo 7, "
    "Şekil 5-6). Bu bulgu, Is_Top3 olasılıklarının sıralama (ranking) amaçlı kullanım "
    "için güvenilir, ancak doğrudan beklenen-değer (EV) hesaplarına sokulmadan önce "
    "yeniden kalibre edilmesi gerektiğini ortaya koymaktadır."
)

BULGU_MCNEMAR = (
    "1.830 koşuluk sızıntısız test kümesinde, modelin doğal seçiminin (olasılığa göre "
    "ilk-k) piyasanın doğal seçimine (favori sırası) kıyasla isabet oranı, herhangi bir "
    "ödeme varsayımı yapılmadan karşılaştırılmıştır (Tablo 8). Modelin altı bahis "
    "türünün tamamında piyasayı geçtiği ve farkın McNemar exact testiyle istatistiksel "
    "olarak anlamlı olduğu (p<0.05) görülmüştür; en güçlü sonuç Ganyan'da elde edilmiştir "
    "(model %39.8 - piyasa %31.7, Δ=+8.1 puan, p=1.6×10⁻¹⁶). Yalnızca Ganyan bahsinde "
    "veride gerçek (piyasa-bağımsız) ödeme bilgisi bulunduğundan, bu türde dairesel "
    "olmayan tek güvenilir kasa simülasyonu kurulabilmiştir: 178 pozitif-EV bahisten "
    "%39.9 isabetle, flat strateji 1.000 TL'yi 3.350 TL'ye çıkarmıştır (ROI=+%132.1; "
    "Tablo 9, Şekil 7). Ancak aynı modelin canlı forward-test'inde (18 koşu, gerçek "
    "ödemelerle) Ganyan ROI'si yaklaşık -%36 olarak gözlenmiştir; bu fark, geçmiş OOF "
    "backtest'inin iyimser olabileceğini ve gerçek hakemin forward-test olduğunu "
    "göstermektedir. Egzotik bahis türleri (İkili, Sıralı İkili, Üçlü, Tabela) için "
    "geçmiş ödeme verisi bulunmadığından bu türlerin ROI'si yalnızca piyasa-ima "
    "yöntemiyle gösterge niteliğinde hesaplanabilmiş ve dairesel olduğu için literal TL "
    "değeri olarak yorumlanmamalıdır; bu türler için de McNemar sonucu (Tablo 8) tek "
    "dairesel-olmayan kanıttır."
)

OZELLIK_ONEMI = (
    "Ağaç bazlı algoritmaların ayırt etme metrikleri (Native Gain ve SHAP analizi) "
    "incelendiğinde; atın bağımsız saf handikap puanından öte rakiplerine olan göreceli "
    "puanı (Relative_Handikap) ve görece sıkleti (Relative_Siklet) ilk sıralarda yer "
    "almıştır. Ayrıca atın son 400 metre idman derecesi, her iki tahmin modelinde de ilk "
    "10 özellik arasına girmiş; bu da idman verisi gibi biyolojik/performans "
    "göstergelerinin isabet oranında belirgin bir kırılım yarattığını göstermektedir "
    "(Şekil 8-11)."
)

SONUC = (
    "Bu çalışmada, at performanslarının tahmini amacıyla kullanılan kolektif öğrenme "
    "modellerinin, sızıntısız (TimeSeriesSplit) bir metodolojiyle değerlendirildiğinde "
    "de hipodrom istatistiklerine uyarlanabilir, ticari ve istatistiksel olarak "
    "savunulabilir bir avantaj sunduğu üç ayrı kanıt hattıyla gösterilmiştir: (i) piyasa "
    "sinyali çıkarıldığında AUC'nin yalnızca 0.01-0.02 düşmesi, modelin gücünün büyük "
    "ölçüde kendi özellik mühendisliğinden geldiğini; (ii) McNemar testlerinin altı "
    "bahis türünün tamamında modelin piyasayı istatistiksel olarak anlamlı şekilde "
    "geçtiğini (p<0.05, çoğunlukla p<0.001) kanıtlamıştır. Bununla birlikte kalibrasyon "
    "analizi ve canlı forward-test sonuçları, ham olasılıkların -özellikle Is_Top3 ve "
    "egzotik bahis EV hesaplarında- doğrudan kullanılmadan önce isotonic/Platt "
    "yöntemleriyle yeniden kalibre edilmesi gerektiğini göstermiştir; bu, çalışmanın bir "
    "sonraki adımı olarak planlanmaktadır. Ayrıca kurulan VPS tabanlı canlı zamanlayıcı "
    "(live scheduler) sayesinde günlük forward-test verisi sürekli birikmekte, modelin "
    "gerçek-dünya performansı zamanla daha güvenilir örneklem büyüklüğüyle "
    "değerlendirilebilecektir. Sonuç olarak ML yöntemlerinin, doğruluk payı yüksek ve "
    "istatistiksel olarak doğrulanmış karar destek sistemlerine hizmet edebileceği; "
    "ancak gerçek bahis uygulamalarında yüksek varyans ve kalibrasyon gereksinimleri "
    "nedeniyle temkinli yaklaşılması gerektiği ortaya konmuştur."
)

TESEKKUR = (
    "Projenin gerçekleştirilmesinde veri alt yapısını sunan ve akademik değerlendirme "
    "teşviklerinde bulunan bilim heyetine ile TJK'nın veri tabanlarına teşekkür ederiz."
)

REFERANSLAR = [
    "[1] Cihan, P. (2024). Horse Surgery and Survival Prediction with Artificial "
    "Intelligence Models: Performance Comparison of Original, Imputed, Balanced, and "
    "Feature-Selected Datasets. Kafkas Üniversitesi Veteriner Fakültesi Dergisi, 30(2), "
    "233-241.",
    "[2] Gupta, M., & Singh, L. (2024). Horse Race Results Prediction Using Machine "
    "Learning Algorithms With Feature Selection. International Journal of Intelligent "
    "Systems and Applications in Engineering, 12(2s), 132-139.",
    "[3] Higgins, M. (2018). The Gambler Who Cracked the Horse-Racing Code. Bloomberg "
    "Businessweek.",
    "[4] Oda, D., & Onogi, A. (2024). Assessing the predictability of racing performance "
    "of Thoroughbreds using mixed-effects model. Journal of Animal Breeding and "
    "Genetics, 141(1), 24-32.",
    "[5] Obi, O. C., et al. (2024). Data science in sports analytics: A review of "
    "performance optimization and fan engagement. Journal of Artificial Intelligence "
    "and Data Science Techniques, 1(2).",
    "[6] Tondapu, N. (2024). Efficient Market Dynamics: Unraveling Informational "
    "Efficiency in UK Horse Racing Betting Markets Through Betfair's Time Series "
    "Analysis. arXiv preprint arXiv:2402.02623.",
    "[7] Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. "
    "Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery "
    "and Data Mining (pp. 785-794).",
    "[8] Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., ... & Liu, T. Y. "
    "(2017). LightGBM: A highly efficient gradient boosting decision tree. Advances in "
    "Neural Information Processing Systems, 30.",
    "[9] Harville, D. A. (1973). Assigning probabilities to the outcomes of "
    "multi-entry competitions. Journal of the American Statistical Association, "
    "68(342), 312-316.",
    "[10] Benter, W. (1994). Computer based horse race handicapping and wagering "
    "systems: A report. In W. T. Ziemba, V. S. Lo, & D. B. Hausch (Eds.), Efficiency of "
    "Racetrack Betting Markets (pp. 183-198). Academic Press.",
    "[11] McNemar, Q. (1947). Note on the sampling error of the difference between "
    "correlated proportions or percentages. Psychometrika, 12(2), 153-157.",
    "[12] Niculescu-Mizil, A., & Caruana, R. (2005). Predicting good probabilities with "
    "supervised learning. Proceedings of the 22nd International Conference on Machine "
    "Learning (ICML), 625-632.",
]


# ───────────────────────── helpers ─────────────────────────

def set_run(run, bold=False, size=BODY_SIZE, italic=False, superscript=False):
    run.font.name = FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.font.superscript = superscript


def add_para(doc, text="", bold=False, size=BODY_SIZE,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        set_run(p.add_run(text), bold=bold, size=size, italic=italic)
    return p


def add_heading(doc, text):
    return add_para(doc, text, bold=True, size=Pt(12))


def style_table(tbl):
    tbl.style = "Table Grid"
    for row in tbl.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_text(cell, text, bold=False, size=TABLE_SIZE):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(str(text)), bold=bold, size=size)


def add_caption_table(doc, caption, header, rows):
    ncols = len(header)
    tbl = doc.add_table(rows=2 + len(rows), cols=ncols)
    style_table(tbl)
    cap_cell = tbl.rows[0].cells[0]
    for c in tbl.rows[0].cells[1:]:
        cap_cell = cap_cell.merge(c)
    set_cell_text(cap_cell, caption, bold=True, size=Pt(11))
    for j, h in enumerate(header):
        set_cell_text(tbl.rows[1].cells[j], h, bold=True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            set_cell_text(tbl.rows[2 + i].cells[j], val)
    doc.add_paragraph()
    return tbl


def add_figure(doc, path, caption, width_cm=14):
    if not os.path.isfile(path):
        add_para(doc, f"[Görsel bulunamadı: {os.path.basename(path)}]",
                 italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    add_para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(11), space_after=12)


def pct(x):
    return f"%{x * 100:.1f}"


def signed_pct(x):
    return f"{'+' if x >= 0 else ''}%{x * 100:.1f}"


# ───────────────────────── data ─────────────────────────

def load_model_tables():
    full = pd.read_csv(os.path.join(BASE_DIR, "model_comparison.csv"))
    abl = pd.read_csv(os.path.join(BASE_DIR, "model_comparison_ablation.csv"))

    winner = full[full["Target"] == "Is_Winner"].sort_values("AUC_mean", ascending=False)
    top3 = full[full["Target"] == "Is_Top3"].sort_values("AUC_mean", ascending=False)

    tablo3 = [[r["Model"], f"{r['AUC_mean']:.4f}", pct(r["P@1"]), signed_pct(r["ROI_value"])]
              for _, r in winner.iterrows()]
    tablo4 = [[r["Model"], f"{r['AUC_mean']:.4f}", pct(r["P@1"]), pct(r["P@3"])]
              for _, r in top3.iterrows()]

    def ablation_rows(target):
        f = full[full["Target"] == target].set_index("Model")
        a = abl[abl["Target"] == target].set_index("Model")
        order = full[full["Target"] == target].sort_values(
            "AUC_mean", ascending=False)["Model"].tolist()
        rows, deltas = [], []
        for m in order:
            auc_f, auc_a = f.loc[m, "AUC_mean"], a.loc[m, "AUC_mean"]
            p1_f, p1_a = f.loc[m, "P@1"], a.loc[m, "P@1"]
            d_auc = auc_f - auc_a
            deltas.append(d_auc)
            rows.append([m, f"{auc_f:.4f}", f"{auc_a:.4f}", f"+{d_auc:.4f}",
                         pct(p1_f), pct(p1_a), signed_pct(p1_f - p1_a)])
        avg = sum(deltas) / len(deltas)
        return rows, avg

    tablo5, avg_winner = ablation_rows("Is_Winner")
    tablo6, avg_top3 = ablation_rows("Is_Top3")
    return tablo3, tablo4, tablo5, avg_winner, tablo6, avg_top3


# ───────────────────────── build ─────────────────────────

def build():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    hdr_p = section.header.paragraphs[0]
    hdr_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(hdr_p.add_run("Aibatyr AMANZHOLOV*, Pınar CİHAN**" + " " * 12 + "TNKÜ FBE (2026)"),
            size=Pt(9), italic=True)

    inst_tbl = doc.add_table(rows=1, cols=1)
    cell = inst_tbl.rows[0].cells[0]
    cell.text = ""
    for i, line in enumerate(["Tekirdağ Namık Kemal Üniversitesi", "Fen Bilimleri Enstitüsü",
                               "Tez Makalesi", "http://fbe.nku.edu.tr"]):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(line), bold=(i < 3), size=Pt(11))
    doc.add_paragraph()

    add_para(doc, TITLE, bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Aibatyr Amanzholov"))
    set_run(p.add_run("1"), superscript=True)
    set_run(p.add_run(", "))
    set_run(p.add_run("Pınar Cihan"))
    set_run(p.add_run("2"), superscript=True)

    add_para(doc, "1 e-posta: 1220606648@gmail.com", size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "2 e-posta: pkaya@nku.edu.tr", size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    meta = doc.add_table(rows=4, cols=2)
    style_table(meta)
    set_cell_text(meta.rows[0].cells[0], "Tez Kabulü", bold=True)
    set_cell_text(meta.rows[0].cells[1], "Makale Türü", bold=True)
    set_cell_text(meta.rows[1].cells[0], "06 Ekim 2021")
    set_cell_text(meta.rows[1].cells[1], "Araştırma Makalesi / Research Article")
    set_cell_text(meta.rows[2].cells[0], "Anahtar Kelimeler", bold=True)
    set_cell_text(meta.rows[2].cells[1], "Özet", bold=True)

    kw_cell = meta.rows[3].cells[0]
    kw_cell.text = ""
    for i, kw in enumerate(KEYWORDS):
        p = kw_cell.paragraphs[0] if i == 0 else kw_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run(p.add_run(kw + ("," if i < len(KEYWORDS) - 1 else "")), size=Pt(10))

    abs_cell = meta.rows[3].cells[1]
    abs_cell.text = ""
    abs_p = abs_cell.paragraphs[0]
    abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_run(abs_p.add_run(ABSTRACT), size=Pt(10))
    doc.add_paragraph()

    # ── GİRİŞ ──
    add_heading(doc, "GİRİŞ")
    add_para(doc, GIRIS_INTRO)
    add_heading(doc, "Literatür Özeti")
    add_para(doc, LITERATUR_OZETI)
    add_heading(doc, "Çalışmanın Amacı ve Kapsamı")
    add_para(doc, AMAC_KAPSAM)

    # ── MATERYAL VE YÖNTEM ──
    add_heading(doc, "MATERYAL VE YÖNTEM")
    add_heading(doc, "Veri Seti ve Ön İşleme")
    add_para(doc, VERI_SETI)
    add_caption_table(
        doc, "Tablo 1: TJK At Yarışı Veri Setinin Özellikleri ve Bilgileri",
        ["Özellik", "Tür", "Özellik Bilgisi", "Yön"],
        [
            ["Şehir, Pist, Cinsiyet, Mesafe", "Kategorik",
             "Hedef kodlamalarla (target encoding) dönüştürülmüş yarış dinamikleri", "Giriş"],
            ["Handikap Puanı, Sıklet", "Sayısal",
             "Atın salt kütlesi/puanı ve rakiplerine göre görece (relative) farklılık değerleri", "Giriş"],
            ["Derece 400/600/800m (sn)", "Sayısal",
             "Min: 21.5 - Maks: 30.0 (idman sprint süreleri)", "Giriş"],
            ["Ganyan Sayı & Olasılık", "Sayısal", "Min: 1.05 - Maks: 150.0 (Piyasa tahmini katsayısı)", "Giriş"],
            ["Jokey, Antrenör, Baba/Anne Başarı", "Sayısal",
             "0.0 ile 1.0 arası hedefe yönelik kümülatif kazanma/tabela oranı", "Giriş"],
            ["Is_Winner, Is_Top3", "Kategorik", "Hayır: 0 / Evet: 1 (Lojistik regresyon sınıfları)", "Çıktı"],
        ],
    )

    add_heading(doc, "Sızıntısız Değerlendirme Metodolojisi")
    add_para(doc, TS_SPLIT)

    add_heading(doc, "Model Algoritmaları ve Yöntemler")
    add_para(doc, MODEL_ALGO)
    add_para(doc, MODEL_ENSEMBLE_DETAY)
    add_caption_table(
        doc, "Tablo 2: Bu çalışmada kullanılan AI (Makine Öğrenmesi) modelleri",
        ["Kategoriler", "Model / Kısaltma", "Açıklama"],
        [
            ["Temel Sınıflandırıcı", "Lojistik Regresyon (LR)", "Verileri uygun bir mantıksal sigmoid eğrisine uyduran temel modeldir."],
            ["Bagging (Torbalama)", "Rastgele Orman (RF)", "Varyansı ve overfitting riskini düşürmek için birden fazla paralel karar ağacını birleştirir."],
            ["Bagging (Torbalama)", "Bagging LGBM", "Sınıflandırıcı olarak ayrı LGBM modellerini farklı veri alt kümelerinde eğitir."],
            ["Boosting (Arttırma)", "Gradient Boosting (GBM)", "Ağaçların hatalarından ders alarak ardışık şekilde birbirini geliştirdiği temel boosting modelidir."],
            ["Boosting (Arttırma)", "XGBoost (XGB)", "Modellerin ardışık olarak güçlendiği, L1/L2 regülarizasyonlu optimize mimaridir."],
            ["Boosting (Arttırma)", "LightGBM (LGB)", "Ağaçların derinlikten çok yaprak bazlı (leaf-wise) büyümesini esas alan çok hızlı yaklaşımdır."],
            ["Boosting (Arttırma)", "CatBoost (CAT)", "Özellikle kategorik özellikler için özel hedef kodlama mekanizmaları barındıran boosting modelidir."],
            ["Voting (Oylama)", "Voting Ensemble", "Bağımsız algoritmaların tahmin olasılıklarını (soft-voting) ortak bir kararda harmanlar."],
            ["Stacking (Yığınlama)", "Stacking Ensemble", "Diğer zayıf öğrenicilerin tahmin olasılıklarını meta-özellik olarak alıp Lojistik meta-modeline sunan birleştiricidir."],
        ],
    )

    add_heading(doc, "Piyasa Sinyali Ablasyon Tasarımı")
    add_para(doc, ABLASYON_YONTEM)
    add_heading(doc, "Olasılık Kalibrasyon Değerlendirmesi")
    add_para(doc, KALIBRASYON_YONTEM)
    add_heading(doc, "Bahis Stratejisi ve Anlamlılık Testi Yöntemi")
    add_para(doc, BAHIS_YONTEM)

    # ── BULGULAR ──
    add_heading(doc, "BULGULAR")
    add_para(doc, ("Eğitilen modellerin performansı \"Birinci Tahmini (Is_Winner)\" ve "
                   "\"Tabela Tahmini (Is_Top3)\" olmak üzere iki ayrı görevde "
                   "ölçümlenmiştir. Başarıları daha salt ve ispatlanabilir kılmak adına, "
                   "her yarışta sadece ganyan favorisi olan ata oynanması durumu referans "
                   "noktası (Ganyan Baseline) alınmıştır."))

    tablo3, tablo4, tablo5, avg_w, tablo6, avg_t = load_model_tables()

    add_heading(doc, "Birinci Tahminin Sonuçları (Is_Winner)")
    add_para(doc, BULGU_WINNER)
    add_caption_table(
        doc, "Tablo 3: Kazanan Tahmini İçin Model Performansları (Is_Winner)",
        ["Model Algoritması", "Eğri Altı Alan (AUC)", "Sıra-1 İsabet Doğruluğu (P@1)", "Değer Bahsi Stratejik Getirisi (ROI)"],
        tablo3,
    )
    add_figure(doc, os.path.join(BASE_DIR, "academic_plot_is_winner.png"),
               "Şekil 1: Birincilik (Is_Winner) tahmini algoritmalarının AUC, isabet oranı (P@1) ve ROI düzeyindeki performans grafikleri")

    add_heading(doc, "Tabela Tahmini Sonuçları (Is_Top3)")
    add_para(doc, BULGU_TOP3)
    add_caption_table(
        doc, "Tablo 4: Tabela Tahmini İçin Model Performansları (Is_Top3)",
        ["Model Algoritması", "Eğri Altı Alan (AUC)", "Sıra-3 İsabet Doğruluğu (P@1)", "Precision@3"],
        tablo4,
    )
    add_figure(doc, os.path.join(BASE_DIR, "academic_plot_is_top3.png"),
               "Şekil 2: Tabela (Is_Top3) tahmini modellerinin performans karşılaştırma grafikleri")

    add_heading(doc, "Piyasa Sinyali Ablasyonu")
    add_para(doc, BULGU_ABLASYON)
    add_caption_table(
        doc, "Tablo 5: Piyasa Sinyali Ablasyonu — Kazanan Tahmini (Is_Winner)",
        ["Model", "AUC (tam)", "AUC (Ganyansız)", "ΔAUC", "P@1 (tam)", "P@1 (Ganyansız)", "ΔP@1"],
        tablo5,
    )
    add_para(doc, f"Ortalama ΔAUC (Is_Winner): +{avg_w:.4f} (piyasa sinyalinin ortalama katkısı).",
              italic=True, size=Pt(10), space_after=10)
    add_caption_table(
        doc, "Tablo 6: Piyasa Sinyali Ablasyonu — Tabela Tahmini (Is_Top3)",
        ["Model", "AUC (tam)", "AUC (Ganyansız)", "ΔAUC", "P@1 (tam)", "P@1 (Ganyansız)", "ΔP@1"],
        tablo6,
    )
    add_para(doc, f"Ortalama ΔAUC (Is_Top3): +{avg_t:.4f} (piyasa sinyalinin ortalama katkısı).",
              italic=True, size=Pt(10), space_after=10)
    add_figure(doc, os.path.join(BASE_DIR, "ablation_auc_is_winner.png"),
               "Şekil 3: Tam model ile piyasa sinyalsiz (Ganyansız) modelin AUC karşılaştırması — Is_Winner")
    add_figure(doc, os.path.join(BASE_DIR, "ablation_auc_is_top3.png"),
               "Şekil 4: Tam model ile piyasa sinyalsiz (Ganyansız) modelin AUC karşılaştırması — Is_Top3")

    add_heading(doc, "Olasılık Kalibrasyonu")
    add_para(doc, BULGU_KALIBRASYON)
    add_caption_table(
        doc, "Tablo 7: Olasılık Kalibrasyonu (OOF, 17.923 at-koşu kaydı)",
        ["Hedef", "Brier Skoru", "ECE", "Taban Oran", "Yorum"],
        [
            ["Is_Winner", "0.0761", "0.0054", "%10.2", "Neredeyse kusursuz kalibre"],
            ["Is_Top3", "0.2630", "0.3018", "%30.7", "Aşırı güvenli (over-confident)"],
        ],
    )
    add_figure(doc, os.path.join(BASE_DIR, "calibration_is_winner.png"),
               "Şekil 5: Güvenilirlik diyagramı (reliability diagram) — Is_Winner")
    add_figure(doc, os.path.join(BASE_DIR, "calibration_is_top3.png"),
               "Şekil 6: Güvenilirlik diyagramı (reliability diagram) — Is_Top3")

    add_heading(doc, "Model vs Piyasa — Bahis Kenarı (McNemar) ve Ganyan Kasası")
    add_para(doc, BULGU_MCNEMAR)
    add_caption_table(
        doc, "Tablo 8: Model vs Piyasa Seçim İsabeti — McNemar Exact Testi (n=1.830 koşu)",
        ["Bahis Türü", "Model İsabet", "Piyasa İsabet", "Δ (puan)", "McNemar p"],
        [
            ["Ganyan", "%39.8", "%31.7", "+8.1", "1.6×10⁻¹⁶"],
            ["İkili", "%22.1", "%16.4", "+5.7", "1.1×10⁻¹⁰"],
            ["Sıralı İkili", "%14.6", "%9.6", "+5.0", "2.5×10⁻¹⁰"],
            ["Plase", "%73.0", "%68.3", "+4.7", "1.9×10⁻⁶"],
            ["Üçlü", "%4.0", "%2.7", "+1.3", "1.4×10⁻²"],
            ["Tabela", "%1.6", "%0.9", "+0.7", "4.1×10⁻²"],
        ],
    )
    add_caption_table(
        doc, "Tablo 9: Ganyan Kasası — Gerçek Oranla Backtest vs Forward-Test",
        ["Ölçüm", "Backtest (OOF, 178 bahis)", "Canlı Forward-Test (18 koşu)"],
        [
            ["İsabet Oranı", "%39.9", "—"],
            ["Flat Kasa ROI", "+%132.1 (1.000→3.350 TL)", "~ −%36"],
            ["Yorum", "İyimser — dairesel değil ama geçmişe dayalı", "Gerçek hakem"],
        ],
    )
    add_figure(doc, os.path.join(BASE_DIR, "bankroll_curve.png"),
               "Şekil 7: Flat ve ¼-Kelly kasa stratejilerinin zaman içindeki bakiye eğrisi (Ganyan bahisleri, gerçek oranlarla, OOF backtest)")

    add_heading(doc, "Özellik Önemi (Feature Importance) ve Değişkenlerin Rolü")
    add_para(doc, OZELLIK_ONEMI)
    add_figure(doc, os.path.join(BASE_DIR, "fi_LightGBM_Is_Winner_native.png"),
               "Şekil 8: LightGBM özellik önemi (native gain) — Is_Winner", width_cm=12)
    add_figure(doc, os.path.join(BASE_DIR, "fi_LightGBM_Is_Winner_shap.png"),
               "Şekil 9: SHAP değerleri — Is_Winner", width_cm=12)
    add_figure(doc, os.path.join(BASE_DIR, "fi_LightGBM_Is_Top3_native.png"),
               "Şekil 10: LightGBM özellik önemi (native gain) — Is_Top3", width_cm=12)
    add_figure(doc, os.path.join(BASE_DIR, "fi_LightGBM_Is_Top3_shap.png"),
               "Şekil 11: SHAP değerleri — Is_Top3", width_cm=12)

    # ── SONUÇ ──
    add_heading(doc, "SONUÇ VE ÖNERİLER")
    add_para(doc, SONUC)

    add_heading(doc, "Teşekkür")
    add_para(doc, TESEKKUR)

    add_heading(doc, "KAYNAKLAR (APA)")
    for ref in REFERANSLAR:
        add_para(doc, ref, size=Pt(11), space_after=4)

    doc.save(OUT_PATH)
    print(f"✓ Oluşturuldu: {OUT_PATH}")


if __name__ == "__main__":
    build()
